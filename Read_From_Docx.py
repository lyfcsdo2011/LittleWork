import docx
import os
import xlwt

# 创建新的Excel
excel = xlwt.Workbook(encoding='utf-8')
sheet = excel.add_sheet('Mysheet')
sheet.write(0, 0, label='姓名')
sheet.write(0, 1, label='学院')
style_num = [42, 43]                                    # 单元格颜色填充色
style_count = 0
pattern1 = xlwt.Pattern()                                # 创建单元格背景
pattern2 = xlwt.Pattern()
pattern1.pattern = xlwt.Pattern.SOLID_PATTERN
pattern1.pattern_fore_colour = style_num[0]
pattern2.pattern = xlwt.Pattern.SOLID_PATTERN
pattern2.pattern_fore_colour = style_num[1]
pattern = [pattern1, pattern2]

style = xlwt.XFStyle()
style.pattern = pattern[style_count]

# 读取目标路径下所有docx文件中表的姓名和学号
path = "./Docx/"
file_name = os.listdir(path)
tab_name_dict = []                                      # 创建姓名集合用于存放姓名
tab_inst_dict = []                                      # 创建学院集合用于存放学院
for i in file_name:
    name = path + i
    doc = docx.Document(name)                           # 打开word文档
    tab_name = doc.tables[0].rows[0].cells[1]           # word中第一个表第一行第二列为姓名
    tab_institute = doc.tables[0].rows[1].cells[3]      # word中第一个表第二行第四列为学院
    # 存储姓名和学院
    tab_name_dict.append(tab_name.text)
    tab_inst_dict.append(tab_institute.text)
    print(tab_name.text+" "+i)
'''
# 开始写入Excel
rows = 1
cols = 0
length = len(tab_name_dict)

for i in range(0, length):
    sheet.write(rows, cols, tab_name_dict[i], style)
    sheet.write(rows, cols + 1, tab_inst_dict[i], style)
    if (i + 1) % 10 == 0:                                # 填充色换色
        style_count = (style_count + 1) % 2
        style.pattern = pattern[style_count]
    if (i + 1) % 90 == 0:                                # 换到另一列
        rows = 1
        cols = cols + 2
        if i != length - 1:
            sheet.write(0, cols, label='姓名')
            sheet.write(0, cols + 1, label='学院')
    else:
        rows = rows + 1

excel.save('3.xls')'''